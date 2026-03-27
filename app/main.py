from __future__ import annotations

import json
import os
import re
from pathlib import Path
from typing import Any, Dict
from uuid import uuid4

from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from google import genai
from google.genai import types
from docx import Document
import pdfplumber

from app.evaluation import evaluate
from app.input_handler import build_config
from app.models import EngineStatus, RecommendRequest, SimulateRequest, SimulateResponse, PRDParseResponse
from app.orchestrator import run_simulation
from app.recommendation import recommend
from app.key_manager import key_manager

app = FastAPI(title="Product Strategy Simulator", version="1.0.0")
ROOT_DIR = Path(__file__).resolve().parents[1]
WARROOM_HTML = ROOT_DIR / "warroom-ui.html"

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def _fallback_parse_prd_text(text_content: str) -> Dict[str, Any]:
    lines = [ln.strip() for ln in text_content.splitlines() if ln.strip()]
    lower = text_content.lower()

    # Product name: prefer explicit header, otherwise first non-empty line.
    product_name = "Untitled Product"
    name_match = re.search(r"(?:product\s*name|name)\s*[:\-]\s*(.+)", text_content, re.IGNORECASE)
    if name_match:
        product_name = name_match.group(1).strip()
    elif lines:
        product_name = lines[0][:80]

    # Description: first paragraph-like sentence block after explicit label.
    desc_match = re.search(
        r"(?:product\s*description|description|overview|summary)\s*[:\-]\s*(.+)",
        text_content,
        re.IGNORECASE,
    )
    if desc_match:
        product_description = desc_match.group(1).strip()[:220]
    else:
        product_description = " ".join(lines[1:4])[:220] if len(lines) > 1 else "No description available"

    # Features: bullet extraction first, then keyword fallback.
    feature_lines = re.findall(r"(?:^|\n)\s*[-*•]\s+([^\n]+)", text_content)
    features = [f.strip()[:60] for f in feature_lines if f.strip()]
    if not features:
        keyword_map = {
            "smart": "Smart Control",
            "portable": "Portability",
            "timer": "Smart Timer",
            "auto": "Auto Automation",
            "energy": "Energy Efficient",
            "app": "App Integration",
        }
        features = [val for key, val in keyword_map.items() if key in lower]
    if not features:
        features = ["Core capability"]

    # Price extraction (supports INR and plain numbers).
    price = None
    price_match = re.search(r"(?:₹|INR\s*)(\d[\d,]*)", text_content, re.IGNORECASE)
    if not price_match:
        price_match = re.search(r"(?:price|base\s*price)\s*[:\-]\s*(\d[\d,]*)", text_content, re.IGNORECASE)
    if price_match:
        try:
            price = float(price_match.group(1).replace(",", ""))
        except ValueError:
            price = None

    pricing_strategy = "competitive"
    if "penetration" in lower:
        pricing_strategy = "penetration"
    elif "premium" in lower:
        pricing_strategy = "premium"
    elif "skimming" in lower:
        pricing_strategy = "skimming"

    target_audience = "middle_income"
    if any(k in lower for k in ["budget", "price-sensitive", "price sensitive"]):
        target_audience = "budget"
    elif any(k in lower for k in ["enterprise", "b2b"]):
        target_audience = "enterprise"
    elif "premium" in lower:
        target_audience = "premium"

    region = "IN"
    if any(k in lower for k in ["united states", "usa", " us "]):
        region = "US"
    elif any(k in lower for k in ["europe", "eu"]):
        region = "EU"
    elif any(k in lower for k in ["south east asia", "southeast asia", "se asia"]):
        region = "SE_ASIA"
    elif any(k in lower for k in ["mena", "middle east", "africa"]):
        region = "MENA"

    market_scenario = "Standard market conditions"
    scenario_match = re.search(r"(?:market\s*scenario|market\s*context|scenario)\s*[:\-]\s*(.+)", text_content, re.IGNORECASE)
    if scenario_match:
        market_scenario = scenario_match.group(1).strip()[:400]
    elif len(lines) > 3:
        market_scenario = " ".join(lines[3:8])[:400]

    return {
        "product_name": product_name,
        "product_description": product_description,
        "features": features[:10],
        "price": price,
        "pricing_strategy": pricing_strategy,
        "target_audience": target_audience,
        "market_scenario": market_scenario,
        "region": region,
        "extraction_confidence": 0.45,
        "warnings": ["Gemini unavailable. Used heuristic parser fallback."],
    }


@app.get("/health", response_model=EngineStatus)
def health() -> EngineStatus:
    return EngineStatus(status="ok")


@app.get("/")
def warroom_dashboard() -> FileResponse:
    if not WARROOM_HTML.exists():
        raise HTTPException(status_code=404, detail="warroom_ui_not_found")
    return FileResponse(str(WARROOM_HTML))


@app.post("/prd/parse", response_model=PRDParseResponse)
async def parse_prd(file: UploadFile = File(...)) -> PRDParseResponse:
    """Parse PRD file (PDF, DOC, DOCX, TXT, MD) and extract product information using Gemini."""
    try:
        # Read file content
        content = await file.read()
        filename = file.filename.lower()
        
        if not content:
            raise HTTPException(status_code=400, detail="File is empty.")
        
        # Extract text based on file type
        text_content = ""
        
        if filename.endswith('.pdf'):
            try:
                import io
                pdf_file = io.BytesIO(content)
                with pdfplumber.open(pdf_file) as pdf:
                    for page in pdf.pages:
                        text_content += page.extract_text() or ""
                        text_content += "\n"
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Failed to parse PDF: {str(e)}") from e
        
        elif filename.endswith(('.doc', '.docx')):
            try:
                import io
                doc_file = io.BytesIO(content)
                doc = Document(doc_file)
                for paragraph in doc.paragraphs:
                    text_content += paragraph.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text_content += cell.text + " | "
                    text_content += "\n"
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Failed to parse DOC/DOCX: {str(e)}") from e
        
        elif filename.endswith(('.txt', '.md')):
            try:
                text_content = content.decode('utf-8')
            except UnicodeDecodeError:
                raise HTTPException(status_code=400, detail="Invalid file encoding. Please use UTF-8.")
        
        else:
            raise HTTPException(
                status_code=400,
                detail="Unsupported file type. Please use: PDF, DOC, DOCX, TXT, or MD"
            )
        
        if not text_content.strip():
            raise HTTPException(status_code=400, detail="No text content could be extracted from file.")
        
        # Truncate if too large (token limits)
        max_length = 20000
        if len(text_content) > max_length:
            text_content = text_content[:max_length]
        
        extraction_prompt = f"""You are a product business analyst. Extract structured information from this PRD (Product Requirements Document).

PRD Content:
{text_content}

Extract and return ONLY valid JSON (no markdown, no explanation) with these fields:
{{
    "product_name": "string - the product name",
    "product_description": "string - brief product description (max 200 chars)",
    "features": ["array", "of", "key", "features"],
    "price": number or null (if mentioned),
    "pricing_strategy": "one of: penetration, competitive, premium, skimming",
    "target_audience": "one of: budget, middle_income, premium, enterprise",
    "market_scenario": "string - market context and challenges described in PRD",
    "region": "two-letter code: IN, US, EU, SE_ASIA, MENA or similar",
    "extraction_confidence": 0.0-1.0 (how confident in the extraction),
    "warnings": ["any", "important", "caveats"]
}}

If a field is not clear from the PRD, use sensible defaults:
- Default pricing_strategy: "competitive"
- Default target_audience: "middle_income"
- Default region: "IN"
- Default price: null (will be set manually)

Return ONLY the JSON object, nothing else."""

        # Use Gemini to parse PRD with Key Rotation Failover
        model_names = ["gemini-1.5-flash", "gemini-2.0-flash"]
        response = None
        last_error: Exception | None = None
        
        # We try each key in our pool if we hit a quota or 429
        available_keys = key_manager.get_all_keys()
        for _ in range(len(available_keys)):
            current_api_key = key_manager.current_key
            if not current_api_key:
                break
                
            client = genai.Client(api_key=current_api_key)
            
            # Try available models for this key
            for model_name in model_names:
                try:
                    response = client.models.generate_content(
                        model=model_name,
                        contents=extraction_prompt,
                        config=types.GenerateContentConfig(
                            temperature=0.7,
                            max_output_tokens=1500,
                        ),
                    )
                    if response:
                        break
                except Exception as exc:
                    last_error = exc
                    msg = str(exc).lower()
                    if "resource_exhausted" in msg or "quota" in msg or "429" in msg:
                        break # break model loop to rotate key
                    
            if response:
                break
            else:
                key_manager.rotate_key()

        if response is None:
            parsed = _fallback_parse_prd_text(text_content)
            parsed["warnings"].append(f"Gemini model unavailable: {last_error}")
        else:
            response_text = response.text

            # Clean up response (remove markdown code blocks if present)
            if response_text.startswith("```"):
                response_text = response_text.split("```")[1]
                if response_text.startswith("json"):
                    response_text = response_text[4:]
            response_text = response_text.strip()

            try:
                parsed = json.loads(response_text)
            except json.JSONDecodeError:
                parsed = _fallback_parse_prd_text(text_content)
                parsed["warnings"].append("Gemini response was not valid JSON. Used heuristic parser fallback.")
        
        # Validate and construct response
        result = PRDParseResponse(
            product_name=parsed.get("product_name", "Untitled Product"),
            product_description=parsed.get("product_description", "No description available"),
            features=parsed.get("features", ["Core capability"]),
            price=parsed.get("price"),
            pricing_strategy=parsed.get("pricing_strategy", "competitive"),
            target_audience=parsed.get("target_audience", "middle_income"),
            market_scenario=parsed.get("market_scenario", "Standard market conditions"),
            region=parsed.get("region", "IN"),
            extraction_confidence=parsed.get("extraction_confidence", 0.8),
            warnings=parsed.get("warnings", []),
        )
        
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PRD parsing error: {str(e)}") from e


@app.post("/simulate", response_model=SimulateResponse)
def simulate(payload: SimulateRequest) -> SimulateResponse:
    try:
        config = build_config(payload.input.model_dump())
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    sim = run_simulation(config, num_steps=payload.num_steps)
    evaluation = evaluate(sim["final_state"], sim["history"])

    market_values: Dict[str, Any] = {
        "fuel_price_index": sim["final_state"]["fuel_price_index"],
        "supply_disruption": sim["final_state"]["supply_disruption"],
        "demand_shift": sim["final_state"]["demand_shift"],
        "market_volatility": sim["final_state"]["market_volatility"],
        "price_sensitivity": sim["final_state"]["price_sensitivity"],
        "market_pressure": sim["final_state"]["market_pressure"],
        "consumer_stress": sim["final_state"]["consumer_stress"],
        "feature_score": sim["final_state"]["feature_score"],
    }

    recommendation = recommend(
        config=config,
        original_score=evaluation["success_score"],
        market_context_values=market_values,
        market_context=sim.get("market_context", {}),
        trends=evaluation.get("trends", {}),
        num_steps=payload.num_steps,
    )

    return SimulateResponse(
        simulation_id=str(uuid4()),
        steps_run=len(sim["history"]),
        history=sim["history"],
        agent_logs=sim["agent_logs"],
        market_context=sim["market_context"],
        evaluation=evaluation,
        recommendation=recommendation,
        final_state=sim["final_state"],
        warnings=config.get("warnings", []) + sim.get("warnings", []),
        diagnostics=sim.get("diagnostics", {}),
    )


@app.post("/pipeline/run")
def run_pipeline(payload: SimulateRequest) -> Dict[str, Any]:
    try:
        config = build_config(payload.input.model_dump())
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    sim = run_simulation(config, num_steps=payload.num_steps)
    evaluation = evaluate(sim["final_state"], sim["history"])

    market_values: Dict[str, Any] = {
        "fuel_price_index": sim["final_state"]["fuel_price_index"],
        "supply_disruption": sim["final_state"]["supply_disruption"],
        "demand_shift": sim["final_state"]["demand_shift"],
        "market_volatility": sim["final_state"]["market_volatility"],
        "price_sensitivity": sim["final_state"]["price_sensitivity"],
        "market_pressure": sim["final_state"]["market_pressure"],
        "consumer_stress": sim["final_state"]["consumer_stress"],
        "feature_score": sim["final_state"]["feature_score"],
    }

    recommendation = recommend(
        config=config,
        original_score=evaluation["success_score"],
        market_context_values=market_values,
        market_context=sim.get("market_context", {}),
        trends=evaluation.get("trends", {}),
        num_steps=payload.num_steps,
    )

    return {
        "pipeline_id": str(uuid4()),
        "input_handler": {
            "normalized_config": config,
            "warnings": config.get("warnings", []),
        },
        "orchestrator": {
            "steps_requested": payload.num_steps,
            "steps_run": len(sim["history"]),
            "warnings": sim.get("warnings", []),
            "diagnostics": sim.get("diagnostics", {}),
        },
        "market_analyzer_agent": {
            "market_context": sim.get("market_context", {}),
            "market_state": {
                "fuel_price_index": sim["final_state"].get("fuel_price_index", 0.0),
                "supply_disruption": sim["final_state"].get("supply_disruption", 0.0),
                "demand_shift": sim["final_state"].get("demand_shift", 0.0),
                "market_volatility": sim["final_state"].get("market_volatility", 0.0),
                "price_sensitivity": sim["final_state"].get("price_sensitivity", 0.0),
                "market_pressure": sim["final_state"].get("market_pressure", 0.0),
                "consumer_stress": sim["final_state"].get("consumer_stress", 0.0),
            },
        },
        "simulation_loop": {
            "history": sim.get("history", []),
            "agent_logs": sim.get("agent_logs", []),
            "final_state": sim.get("final_state", {}),
        },
        "evaluation_engine": evaluation,
        "recommendation_engine": recommendation,
    }


@app.post("/recommend")
def recommend_endpoint(payload: RecommendRequest) -> Dict[str, Any]:
    required_keys = [
        "fuel_price_index",
        "supply_disruption",
        "demand_shift",
        "market_volatility",
        "price_sensitivity",
        "market_pressure",
        "consumer_stress",
        "feature_score",
    ]

    if payload.current_state:
        market_values = {k: payload.current_state.get(k, 0.5) for k in required_keys}
    else:
        market_values = {k: 0.5 for k in required_keys}

    try:
        return recommend(
            config=payload.config,
            original_score=payload.current_score,
            market_context_values=market_values,
            market_context=payload.market_context,
            num_steps=payload.num_steps,
        )
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"recommendation_failed:{exc}") from exc


@app.post("/diagnostics/simulate")
def simulate_diagnostics(payload: SimulateRequest) -> Dict[str, Any]:
    try:
        config = build_config(payload.input.model_dump())
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    sim = run_simulation(config, num_steps=payload.num_steps)
    evaluation = evaluate(sim["final_state"], sim["history"])

    return {
        "simulation_id": str(uuid4()),
        "steps_run": len(sim["history"]),
        "market_context": sim.get("market_context", {}),
        "warnings": config.get("warnings", []) + sim.get("warnings", []),
        "diagnostics": sim.get("diagnostics", {}),
        "evaluation": {
            "success_score": evaluation.get("success_score"),
            "status": evaluation.get("status"),
            "trends": evaluation.get("trends", {}),
        },
    }
